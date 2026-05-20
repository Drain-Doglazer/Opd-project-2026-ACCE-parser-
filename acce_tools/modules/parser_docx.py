"""parser_docx.py — Reads a DOCX equipment list and returns raw ACCERecord objects.

Single responsibility: open the file, extract table rows, map cells to ACCERecord fields.
Classification (acce_item_symbol / acce_item_type) is NOT performed here.
"""

import re
from typing import Optional

from docx import Document

from .schema import ACCERecord


# ══════════════════════════════════════════════════════════════════
# TRANSLITERATION — Cyrillic tags → Latin for ACCE user_tag
# ══════════════════════════════════════════════════════════════════

_CYR_TO_LAT = {
    'А':'A','Б':'B','В':'V','Г':'G','Д':'D','Е':'E','Ё':'Yo',
    'Ж':'Zh','З':'Z','И':'I','Й':'Y','К':'K','Л':'L','М':'M',
    'Н':'N','О':'O','П':'P','Р':'R','С':'S','Т':'T','У':'U',
    'Ф':'F','Х':'Kh','Ц':'Ts','Ч':'Ch','Ш':'Sh','Щ':'Sch',
    'Ъ':'','Ы':'Y','Ь':'','Э':'E','Ю':'Yu','Я':'Ya',
    'а':'a','б':'b','в':'v','г':'g','д':'d','е':'e','ё':'yo',
    'ж':'zh','з':'z','и':'i','й':'y','к':'k','л':'l','м':'m',
    'н':'n','о':'o','п':'p','р':'r','с':'s','т':'t','у':'u',
    'ф':'f','х':'kh','ц':'ts','ч':'ch','ш':'sh','щ':'sch',
    'ъ':'','ы':'y','ь':'','э':'e','ю':'yu','я':'ya',
}

_FORBIDDEN = re.compile(r'[!@#$%&*()?/{}[\]*=<>,`~^:;|"\' \\+\s～]')


def _transliterate(text: str) -> str:
    return ''.join(_CYR_TO_LAT.get(ch, ch) for ch in text)


def _make_user_tag(raw_tag: str) -> str:
    """Creates a valid ACCE user_tag from a Russian equipment tag (max 20 chars)."""
    lat = _transliterate(raw_tag.strip())
    safe = _FORBIDDEN.sub('-', lat)
    safe = re.sub(r'-{2,}', '-', safe).strip('-')[:20]
    return safe if safe else f"TAG-{raw_tag[:10]}"


# ══════════════════════════════════════════════════════════════════
# TECHNICAL CHARACTERISTICS PARSER (Russian-language patterns)
# ══════════════════════════════════════════════════════════════════

_SEP = r'\s*[–—\-=]\s*'
_NUM = r'([\d\s,\.]+)'
_NUM_MINUS = r'((?:минус\s+)?[\d\s,\.]+)'


def _num(s: str) -> Optional[float]:
    """Parses a number from a string, including 'минус 40' → -40.0."""
    if s is None:
        return None
    s = s.strip()
    negative = bool(re.match(r'минус', s, re.I))
    digits = re.sub(r'[^\d,\.]', '', s.replace(',', '.'))
    if not digits:
        return None
    try:
        val = float(digits)
        return -val if negative else val
    except ValueError:
        return None


def _parse_tech(text: str) -> dict:
    """Parses a Russian-language technical spec string into structured fields."""
    if not text:
        return {}

    r = {}

    def _find(pattern: str) -> Optional[float]:
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if m:
            return _num(m.group(1))
        return None

    # Design temperature: Трасч.
    t_calc = re.search(
        r'Трасч[\.]*\s*[–—\-=]\s*' + _NUM_MINUS + r'\s*°?С',
        text, re.IGNORECASE)
    if t_calc:
        r['design_temperature'] = _num(t_calc.group(1))

    # Operating temperature: Траб.
    t_work = re.search(
        r'Траб[\.]*\s*[–—\-=]\s*' + _NUM_MINUS + r'\s*°?С',
        text, re.IGNORECASE)
    if t_work:
        r['operating_temperature'] = _num(t_work.group(1))

    # Pressure (priority: Ррасч > Рнагнетания > нагнетание > Рраб)
    p_calc  = re.search(r'Ррасч[\.]*\s*[–—\-=]\s*' + _NUM + r'\s*МПа', text, re.IGNORECASE)
    p_nagn  = re.search(r'Рнагнетани[яе]\s*[–—\-=]\s*' + _NUM + r'\s*МПа', text, re.IGNORECASE)
    p_nagn2 = re.search(r'нагнетание\s*[–—\-=]\s*' + _NUM + r'\s*МПа', text, re.IGNORECASE)
    p_work  = re.search(r'Рраб[\.]*\s*[–—\-=]\s*' + _NUM + r'\s*МПа', text, re.IGNORECASE)

    if p_calc:
        r['pressure'] = _num(p_calc.group(1))
    elif p_nagn:
        r['pressure'] = _num(p_nagn.group(1))
    elif p_nagn2:
        r['pressure'] = _num(p_nagn2.group(1))
    elif p_work:
        r['pressure'] = _num(p_work.group(1))

    # Volume
    vol = re.search(r'Объем\s*[–—\-=]\s*' + _NUM + r'\s*м3', text, re.IGNORECASE)
    if vol:
        r['volume'] = _num(vol.group(1))

    # Diameter
    d_stv = re.search(r'Диаметр\s+ствола\s*[–—\-=]\s*' + _NUM + r'\s*(мм|м)\b', text, re.IGNORECASE)
    d_gen = re.search(r'Диаметр\s*[–—\-=]\s*' + _NUM + r'\s*(мм|м)\b', text, re.IGNORECASE)
    dm = d_stv or d_gen
    if dm:
        val = _num(dm.group(1))
        unit = dm.group(2).lower()
        if val is not None:
            r['diameter_m'] = val / 1000 if unit == 'мм' else val

    # Length / height
    ln = re.search(
        r'(?:Длина|Высота\s+ствола|Длина\s*\(высота\))\s*[–—\-=]\s*'
        + _NUM + r'\s*(мм|м)\b',
        text, re.IGNORECASE)
    if ln:
        val = _num(ln.group(1))
        unit = ln.group(2).lower()
        if val is not None:
            r['length_m'] = val / 1000 if unit == 'мм' else val

    # Flow rate / throughput
    prod = re.search(
        r'Производительность\s*[–—\-=]\s*' + _NUM
        + r'\s*(нм3/ч|м3/ч|мл/ч|м3/час)',
        text, re.IGNORECASE)
    if prod:
        val = _num(prod.group(1))
        unit = prod.group(2).lower()
        if unit == 'мл/ч' and val is not None:
            val = val / 1_000_000
            unit = 'м3/ч'
        r['flow_rate']      = val
        r['flow_rate_unit'] = unit

    # Motor power
    pw = re.search(r'Мощность\s*[–—\-=]\s*' + _NUM + r'\s*кВт', text, re.IGNORECASE)
    if pw:
        r['motor_power_kw'] = _num(pw.group(1))

    # Heat transfer area
    area = re.search(
        r'Площадь\s+теплообмена\s*[–—\-=]\s*' + _NUM + r'\s*м2',
        text, re.IGNORECASE)
    if area:
        r['heat_transfer_area_m2'] = _num(area.group(1))

    # Heat duty
    heat = re.search(
        r'Тепловая\s+нагрузка\s*[–—\-=]\s*' + _NUM + r'\s*Гкал/ч',
        text, re.IGNORECASE)
    if heat:
        r['heat_duty_gcalh'] = _num(heat.group(1))

    # Mass (tonnes → kg)
    mass = re.search(r'Масса\s*[–—\-=]\s*' + _NUM + r'\s*т\.?', text, re.IGNORECASE)
    if mass:
        val = _num(mass.group(1))
        if val is not None:
            r['weight_unit_kg'] = val * 1000

    return r


# ══════════════════════════════════════════════════════════════════
# TAG SPLITTING — "Р-001 Р-002 Р-003" → ['Р-001', 'Р-002', 'Р-003']
# ══════════════════════════════════════════════════════════════════

def _split_tags(raw: str) -> list[str]:
    """Splits a cell with one or more tags into a list of individual tags."""
    cyr_tags = re.findall(r'[А-ЯЁ]{1,2}[а-яё]?-\d+', raw)
    if cyr_tags:
        return cyr_tags
    lat_tags = re.findall(r'[A-Z]{1,4}-\d+', raw)
    return lat_tags if lat_tags else []


# ══════════════════════════════════════════════════════════════════
# DEFAULT COLUMN INDICES (for Samples_3.docx layout)
# ══════════════════════════════════════════════════════════════════
COL_NUM  = 0
COL_TAG  = 1
COL_NAME = 2
COL_QTY  = 3
COL_MAT  = 5
COL_TECH = 6

HEADER_ROWS = 2


def _detect_cols(table) -> tuple[int, int]:
    """Auto-detects column indices for 'Материал' and 'Техническая характеристика'."""
    col_mat = -1
    col_tech = -1
    for ri in range(min(2, len(table.rows))):
        cells = [c.text.strip().lower() for c in table.rows[ri].cells]
        for ci, h in enumerate(cells):
            if 'материал' in h and col_mat < 0:
                col_mat = ci
            if re.search(r'характеристик|тех\.?\s*х', h) and col_tech < 0:
                col_tech = ci
    return col_mat, col_tech


# ══════════════════════════════════════════════════════════════════
# PUBLIC API
# ══════════════════════════════════════════════════════════════════

def parse(file_path: str) -> list[ACCERecord]:
    """Reads a DOCX equipment list and returns a list of ACCERecord objects."""
    if not __import__('os').path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)

    if not doc.tables:
        raise ValueError(f"No tables found in {file_path}")

    table = doc.tables[0]
    records: list[ACCERecord] = []
    doc_name = file_path.split('/')[-1]

    auto_mat, auto_tech = _detect_cols(table)
    col_mat  = auto_mat  if auto_mat  >= 0 else COL_MAT
    col_tech = auto_tech if auto_tech >= 0 else COL_TECH

    for row_idx, row in enumerate(table.rows):
        if row_idx < HEADER_ROWS:
            continue

        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]

        if not any(cells):
            continue

        raw_tags_cell = cells[COL_TAG]  if len(cells) > COL_TAG  else ""
        name_raw      = cells[COL_NAME] if len(cells) > COL_NAME else ""
        qty_raw       = cells[COL_QTY]  if len(cells) > COL_QTY  else ""
        material      = cells[col_mat]  if col_mat  >= 0 and len(cells) > col_mat  else ""
        tech_raw      = cells[col_tech] if col_tech >= 0 and len(cells) > col_tech else ""

        try:
            qty = int(re.search(r'\d+', qty_raw).group()) if qty_raw else 1
        except (AttributeError, ValueError):
            qty = 1

        tags = _split_tags(raw_tags_cell)

        if not tags:
            continue

        tech = _parse_tech(tech_raw)

        if len(tags) > 1:
            items_to_create = [(tag, 1) for tag in tags]
        elif qty > 1:
            items_to_create = [(tags[0], 1)] * qty
        else:
            items_to_create = [(tags[0], 1)]

        for item_idx, (raw_tag, _) in enumerate(items_to_create):
            base_tag = _make_user_tag(raw_tag)

            if qty > 1 and len(tags) == 1:
                user_tag = f"{base_tag}-{item_idx + 1}"[:20]
            else:
                user_tag = base_tag

            record = ACCERecord(
                source_file  = file_path,
                source_sheet = doc_name,
                source_row   = row_idx,

                raw_tag        = raw_tag,
                user_tag       = user_tag,
                description_ru = name_raw,
                parent_area    = "",
                quantity       = 1,
                material       = material,

                weight_unit_kg = tech.get('weight_unit_kg'),

                tech_raw               = tech_raw,
                design_temperature     = tech.get('design_temperature'),
                operating_temperature  = tech.get('operating_temperature'),
                pressure               = tech.get('pressure'),
                pressure_unit          = "МПа",
                volume                 = tech.get('volume'),
                diameter_m             = tech.get('diameter_m'),
                length_m               = tech.get('length_m'),
                flow_rate              = tech.get('flow_rate'),
                flow_rate_unit         = tech.get('flow_rate_unit', ''),
                motor_power_kw         = tech.get('motor_power_kw'),
                heat_transfer_area_m2  = tech.get('heat_transfer_area_m2'),
                heat_duty_gcalh        = tech.get('heat_duty_gcalh'),

                action   = "NEW",
                is_valid = True,
            )

            records.append(record)

    return records
